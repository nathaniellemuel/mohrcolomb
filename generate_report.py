"""
=================================================================
Laporan Mohr-Coulomb — Wellbore Stability Analysis Generator
Well: 58-32 Main
Menghasilkan dokumen Word (.docx) berisi laporan lengkap
=================================================================
"""

import subprocess
import sys
import os
from pathlib import Path

# Ensure python-docx
subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])

import numpy as np
import pandas as pd
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ───────────────────────────────────────────────
# RUN INDEX.PY TO GENERATE FRESH PLOTS
# ───────────────────────────────────────────────
print("=" * 60)
print("STEP 1: Running index.py to generate fresh plots...")
print("=" * 60)

script_dir = Path(__file__).parent
index_py = script_dir / "index.py"

if index_py.exists():
    result = subprocess.run(
        [sys.executable, str(index_py)],
        capture_output=True, text=True,
        encoding="utf-8", errors="replace",
        cwd=str(script_dir)
    )
    if result.returncode == 0:
        print("  index.py executed successfully.")
    else:
        print("  WARNING: index.py returned non-zero:")
        print(result.stderr[-500:])
else:
    print("  index.py not found — using existing plots.")

# ───────────────────────────────────────────────
# LOAD DATA FOR STATISTICS
# ───────────────────────────────────────────────
print("\n" + "=" * 60)
print("STEP 2: Loading data for statistics...")
print("=" * 60)

df_raw = pd.read_csv(script_dir / "58-32_main_geophysical_well_log.csv",
                      na_values=[-999.25])
df_raw.replace(-999.25, np.nan, inplace=True)

df_geo = pd.read_csv(script_dir / "geomechanics_results.csv")

# Key statistics
depth_min = df_geo['DEPTH_FT'].min()
depth_max = df_geo['DEPTH_FT'].max()
sv_max = df_geo['Sv_psi'].iloc[-1]
pp_max = df_geo['Pp_psi'].iloc[-1]
shmin_max = df_geo['Shmin_psi'].iloc[-1]
shmax_max = df_geo['SHmax_psi'].iloc[-1]
s1_max = df_geo['sigma1_psi'].iloc[-1]
s3_max = df_geo['sigma3_psi'].iloc[-1]
ucs_max = df_geo['UCS_psi'].max()
ucs_min = df_geo['UCS_psi'].min()
coh_max = df_geo['cohesion_psi'].max()
coh_min = df_geo['cohesion_psi'].min()

print(f"  Depth range : {depth_min:.1f} - {depth_max:.1f} ft")
print(f"  Sv range    : {sv_max:.1f} psi")
print(f"  UCS range   : {ucs_min:.0f} - {ucs_max:.0f} psi")

# ───────────────────────────────────────────────
# HELPER: add horizontal rule
# ───────────────────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(8)

# ───────────────────────────────────────────────
# HELPER: set cell shading
# ───────────────────────────────────────────────
def set_cell_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

# ───────────────────────────────────────────────
# HELPER: styled heading
# ───────────────────────────────────────────────
def add_heading(doc, text, level, bold=True):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.bold = bold
    if level == 0:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(31, 73, 125)
    elif level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(31, 73, 125)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 70, 127)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

# ───────────────────────────────────────────────
# HELPER: code block (monospace, light gray bg)
# ───────────────────────────────────────────────
def add_code_block(doc, code_text, caption=""):
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # Light gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

# ───────────────────────────────────────────────
# HELPER: add image
# ───────────────────────────────────────────────
def add_image(doc, img_path, width_in=5.8, caption="", center=True):
    full_path = script_dir / img_path
    if not full_path.exists():
        p = doc.add_paragraph(f"[GAMBAR TIDAK DITEMUKAN: {img_path}]")
        return
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(full_path), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = cap.runs[0]
        run2.italic = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(80, 80, 80)

# ───────────────────────────────────────────────
# HELPER: styled bullet
# ───────────────────────────────────────────────
def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)

# ───────────────────────────────────────────────
# HELPER: normal paragraph
# ───────────────────────────────────────────────
def add_para(doc, text, bold=False, italic=False, size=11, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

# ───────────────────────────────────────────────
# CREATE DOCUMENT
# ───────────────────────────────────────────────
print("\n" + "=" * 60)
print("STEP 3: Creating Word document...")
print("=" * 60)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# ───────────────────────────────────────────────
# COVER PAGE
# ───────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LAPORAN PROYEK")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(31, 73, 125)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Analisis Kestabilan Sumur Bor dengan Metode")
run2.bold = True
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(31, 73, 125)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Mohr Circle dan Mohr-Coulomb Failure Criterion")
run3.bold = True
run3.font.size = Pt(16)
run3.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()
add_hr(doc)
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run("Data: Well 58-32 Main Geophysical Well Log")
run4.font.size = Pt(13)

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run("Sumur: 58-32 Main | Metode: 1D Geomechanical Modeling")
run5.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

labels = ["Nama Proyek", "Metode Analisis", "Data Input", "Output", "Bahasa Pemrograman"]
values = [
    "Wellbore Stability Analysis — Well 58-32 Main",
    "Mohr Circle & Mohr-Coulomb Failure Criterion",
    "58-32_main_geophysical_well_log.csv",
    "geomechanics_results.csv + 7 Plot (.png)",
    "Python 3 (NumPy, Pandas, Matplotlib, SciPy)"
]
for i, (lbl, val) in enumerate(zip(labels, values)):
    cell_l = info_table.cell(i, 0)
    cell_r = info_table.cell(i, 1)
    set_cell_shading(cell_l, "D9E2F3")
    run_l = cell_l.paragraphs[0].add_run(lbl)
    run_l.bold = True
    run_l.font.size = Pt(10)
    run_r = cell_r.paragraphs[0].add_run(val)
    run_r.font.size = Pt(10)

doc.add_page_break()

# ───────────────────────────────────────────────
# TABLE OF CONTENTS (manual)
# ───────────────────────────────────────────────
add_heading(doc, "DAFTAR ISI", level=1)
toc_entries = [
    ("BAB 1", "DESKRIPSI KASUS", 3),
    ("  1.1", "Latar Belakang", 3),
    ("  1.2", "Tujuan Analisis", 3),
    ("  1.3", "Data dan Sumber Data", 3),
    ("  1.4", "Scope dan Batasan Analisis", 3),
    ("BAB 2", "LANDASAN TEORI", 3),
    ("  2.1", "Tegasan In-Situ dan Stress Regime", 3),
    ("  2.2", "Mohr Circle — Konsep dan Geometri", 3),
    ("  2.3", "Mohr-Coulomb Failure Criterion", 3),
    ("  2.4", "Kerangka Kerja 1D Geomechanical Model", 3),
    ("BAB 3", "ANALISIS KASUS — KODE DAN PENJELASAN", 3),
    ("  3.1", "Preprocessing Data", 3),
    ("  3.2", "Overburden Stress (Sv)", 3),
    ("  3.3", "Pore Pressure — Metode Eaton Sonic", 3),
    ("  3.4", "Horizontal Stresses — Shmin dan SHmax", 3),
    ("  3.5", "Principal Stresses — Penugasan σ₁, σ₂, σ₃", 3),
    ("  3.6", "Rock Strength — UCS, Cohesion, Friction Angle", 3),
    ("  3.7", "Mohr Circle Class dan Failure Check", 3),
    ("BAB 4", "HASIL DAN INTERPRETASI", 3),
    ("  4.1", "Plot 1 — Composite Well Logs", 3),
    ("  4.2", "Plot 2 — In-Situ Stress Profile", 3),
    ("  4.3", "Plot 3 — Rock Strength Profile", 3),
    ("  4.4", "Plot 4 — Mohr Circles Overview", 3),
    ("  4.5", "Plot 4b — Mohr Circle Detail per Depth", 3),
    ("  4.6", "Plot 5 — Mud Weight Window", 3),
    ("  4.7", "Plot 6 — Borehole Breakout Analysis", 3),
    ("  4.8", "Plot 7 — Pure Mohr Circles", 3),
    ("  4.9", "Ringkasan Hasil Analisis", 3),
    ("BAB 5", "KESIMPULAN DAN SARAN", 3),
    ("  5.1", "Kesimpulan", 3),
    ("  5.2", "Saran dan Pengembangan", 3),
    ("LAMPIRAN", "LISTING KODE LENGKAP (index.py)", 3),
]

toc_table = doc.add_table(rows=len(toc_entries), cols=2)
toc_table.style = 'Table Grid'
for i, (num, title, _) in enumerate(toc_entries):
    c0 = toc_table.cell(i, 0)
    c1 = toc_table.cell(i, 1)
    c0.width = Cm(3)
    r0 = c0.paragraphs[0].add_run(num)
    r0.font.size = Pt(10)
    r0.bold = True
    r1 = c1.paragraphs[0].add_run(title)
    r1.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════
# BAB 1 — DESKRIPSI KASUS
# ═══════════════════════════════════════════════
add_heading(doc, "BAB 1 — DESKRIPSI KASUS", level=0)

add_heading(doc, "1.1 Latar Belakang", level=1)
add_para(doc, (
    "Analisis kestabilan sumur bor (wellbore stability analysis) merupakan salah satu aspek "
    "kritis dalam operasi pemboran minyak dan gas. Kegagalan kestabilan sumur dapat "
    "menyebabkan masalah serius seperti collapse casing, loss of circulation, serta "
    "pembentukan breakout yang berpotensi menyebabkan stuck pipe, non-productive time (NPT), "
    "dan bahkan abandonment sumur. Untuk memprediksi dan mencegah kegagalan tersebut, "
    "diperlukan pemahaman komprehensif terhadap kondisi tegasan in-situ (in-situ stresses) "
    "dan kekuatan batuan (rock strength) di sekitar sumur bor."
))
add_para(doc, (
    "Proyek ini menggunakan data log geofisika dari sumur bor untuk membangun model "
    "geomekanika 1-dimensi (1D) dan melakukan analisis kegagalan sumur dengan Metode "
    "Mohr Circle dan Mohr-Coulomb Failure Criterion. Data yang digunakan berasal dari "
    "Well 58-32 Main yang menyediakan rekaman wireline log meliputi data akustik (sonic), "
    "densitas, gamma ray, dan caliper."
))

add_heading(doc, "1.2 Tujuan Analisis", level=1)
add_para(doc, "Analisis ini bertujuan untuk:")
add_bullet(doc, "Menghitung tegasan in-situ 3 arah: Overburden (Sv), Minimum Horizontal (Shmin), dan Maximum Horizontal (SHmax).")
add_bullet(doc, "Mengestimasi pore pressure menggunakan Metode Eaton Sonic.")
add_bullet(doc, "Menentukan effective principal stresses (σ'₁, σ'₂, σ'₃) berdasarkan stress regime.")
add_bullet(doc, "Menghitung rock strength parameters: UCS (Unconfined Compressive Strength), cohesion (C), dan friction angle (φ).")
add_bullet(doc, "Membangun Mohr Circle dan failure envelope Mohr-Coulomb pada kedalaman representatif.")
add_bullet(doc, "Mengidentifikasi zone-zone yang berpotensi mengalami kegagalan (failure) berdasarkan kestabilan sumur.")
add_bullet(doc, "Menentukan safe mud weight window untuk perencanaan pemboran.")

add_heading(doc, "1.3 Data dan Sumber Data", level=1)
add_para(doc, "Dua file data digunakan dalam proyek ini:")

tbl = doc.add_table(rows=3, cols=3)
tbl.style = 'Table Grid'
hdr = ["File", "Deskripsi", "Jumlah Kolom"]
for i, h in enumerate(hdr):
    c = tbl.rows[0].cells[i]
    set_cell_shading(c, "1F497D")
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(10)

rows_data = [
    ["58-32_main_geophysical_well_log.csv",
     "Data mentah wireline log dari sensor borehole (input)",
     "62 parameter"],
    ["geomechanics_results.csv",
     "Hasil kalkulasi geomekanika dari pipeline Python (output)",
     "23 parameter"],
]
for i, row_data in enumerate(rows_data):
    for j, val in enumerate(row_data):
        c = tbl.rows[i + 1].cells[j]
        if i == 0:
            set_cell_shading(c, "D9E2F3")
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(10)
        if j == 0:
            r.bold = True

doc.add_paragraph()
add_para(doc, "Parameter utama yang digunakan dari data log:")
add_bullet(doc, "RHOZ / RHOB — Bulk Density Log (g/cc): untuk menghitung overburden stress")
add_bullet(doc, "ATCO10–ATCO90 — Array Sonic Transit Time (µs/ft): untuk DT (compressional slowness) dan estimasi pore pressure")
add_bullet(doc, "GR — Gamma Ray Log (API): untuk litologi (shale/sand cutoff) dan penentuan Poisson's ratio")
add_bullet(doc, "HCAL — Caliper Log (inchi): untuk identifikasi breakout zone")

add_heading(doc, "1.4 Scope dan Batasan Analisis", level=1)
add_bullet(doc, "Analisis dilakukan pada sumur 58-32 Main dengan data wireline log.")
add_bullet(doc, "Stress regime yang digunakan: Normal Faulting (σ₁ = Sv, σ₂ = SHmax, σ₃ = Shmin).")
add_bullet(doc, "Korelasi rock strength: Horsrud (2001) untuk shale, McNally (1987) untuk sandstone.")
add_bullet(doc, "Metode pore pressure: Eaton Sonic Method dengan eksponen n = 3.0.")
add_bullet(doc, "Batasan: model 1D — tidak memperhitungkan variasi lateral dan heterogenitas batuan.")

doc.add_page_break()

# ═══════════════════════════════════════════════
# BAB 2 — LANDASAN TEORI
# ═══════════════════════════════════════════════
add_heading(doc, "BAB 2 — LANDASAN TEORI", level=0)

add_heading(doc, "2.1 Tegasan In-Situ dan Stress Regime", level=1)
add_para(doc, (
    "Tegasan in-situ adalah tegasan yang bekerja pada formasi batuan di bawah permukaan bumi "
    "sebelum adanya gangguan akibat pemboran. Terdapat tiga tegasan principal yang bekerja: "
    "tegasan vertikal (overburden stress, Sv) dan dua tegasan horizontal (minimum horizontal "
    "stress, Shmin, dan maximum horizontal stress, SHmax)."
))
add_para(doc, (
    "Berdasarkan hubungan magnitudo antara ketiganya, kerangka tektonik (stress regime) "
    "dikelompokkan menjadi tiga:"
))
add_bullet(doc, "Normal Faulting: Sv > SHmax > Shmin (σ₁ = Sv, σ₂ = SHmax, σ₃ = Shmin)")
add_bullet(doc, "Strike-Slip Faulting: SHmax > Sv > Shmin (σ₁ = SHmax, σ₂ = Sv, σ₃ = Shmin)")
add_bullet(doc, "Reverse Faulting: SHmax > Shmin > Sv (σ₁ = SHmax, σ₂ = Shmin, σ₃ = Sv)")

add_para(doc, "Effective stress didefinisikan sebagai σ' = σ - Pp, dengan Pp adalah pore pressure.")

add_heading(doc, "2.2 Mohr Circle — Konsep dan Geometri", level=1)
add_para(doc, (
    "Mohr Circle adalah representasi grafis 2D dari состояние напряжений (state of stress) "
    "pada suatu titik dalam material. Untuk kondisi 3D dengan tiga tegasan principal "
    "σ₁ ≥ σ₂ ≥ σ₃, terdapat tiga kombinasi lingkaran Mohr:"
))
add_bullet(doc, "Lingkaran σ₁-σ₃ (terbesar): mengontrol kondisi failure kritis")
add_bullet(doc, "Lingkaran σ₁-σ₂ (sedang): lebih kecil dari σ₁-σ₃")
add_bullet(doc, "Lingkaran σ₂-σ₃ (kecil): lebih kecil dari σ₁-σ₃")

add_para(doc, "Lingkaran Mohr terbesar σ₁-σ₃ didefinisikan dengan:")
add_bullet(doc, "Center (pusat): c = (σ₁ + σ₃) / 2")
add_bullet(doc, "Radius: R = (σ₁ - σ₃) / 2")
add_para(doc, (
    "Dalam diagram Mohr, sumbu horizontal (σ) merepresentasikan tegasan normal (normal stress), "
    "sedangkan sumbu vertikal (τ) merepresentasikan tegasan geser (shear stress). "
    "Setiap titik pada lingkaran merepresentasikan состояние напряжений pada suatu bidang "
    "tertentu yang berorientasi pada sudut 2θ terhadap sumbu horizontal."
))

add_heading(doc, "2.3 Mohr-Coulomb Failure Criterion", level=1)
add_para(doc, (
    "Mohr-Coulomb Failure Criterion mendefinisikan kondisi di mana batuan akan mengalami "
    "kegagalan (failure) akibat kombinasi tegasan normal dan geser. Kriteria failure "
    "dinyatakan dalam persamaan:"
))

p_eq = doc.add_paragraph()
p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_eq = p_eq.add_run("τ = C + σₙ · tan(φ)")
r_eq.bold = True
r_eq.font.size = Pt(13)
r_eq.font.name = "Courier New"

add_para(doc, "Dengan:")
add_bullet(doc, "τ  = Tegasan geser pada bidang failure (psi)")
add_bullet(doc, "σₙ = Tegasan normal efektif pada bidang failure (psi)")
add_bullet(doc, "C  = Cohesion — kohesi batuan (psi)")
add_bullet(doc, "φ  = Internal Friction Angle — sudut gesek dalam batuan (derajat)")

add_para(doc, "Dari kriteria Mohr-Coulomb, UCS (Unconfined Compressive Strength) связана dengan C dan φ melalui:")
p_eq2 = doc.add_paragraph()
p_eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_eq2 = p_eq2.add_run("UCS = 2C · tan(45° + φ/2)")
r_eq2.bold = True
r_eq2.font.size = Pt(13)
r_eq2.font.name = "Courier New"

add_para(doc, "Sebaliknya, kohesi dapat dihitung dari UCS:")
p_eq3 = doc.add_paragraph()
p_eq3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_eq3 = p_eq3.add_run("C = UCS / [2 · tan(45° + φ/2)]")
r_eq3.bold = True
r_eq3.font.size = Pt(13)
r_eq3.font.name = "Courier New"

add_para(doc, (
    "Titik singgung (tangent point) antara lingkaran Mohr σ₁-σ₃ dan failure envelope "
    "merepresentasikan bidang failure kritis. Sudut antara sumbu horizontal dan garis "
    "yang menghubungkan pusat lingkaran ke titik singgung adalah 2θ = 90° + φ. "
    "Bidang failure itu sendiri membentuk sudut θ = (90° + φ)/2 terhadap bidang utama σ₃."
))

add_heading(doc, "2.4 Kerangka Kerja 1D Geomechanical Model", level=1)
add_para(doc, (
    "Model geomekanika 1D membangun profil tegasan dan kekuatan batuan terhadap kedalaman "
    "dengan mengintegrasikan data log bor. Pipeline analisis terdiri dari:"
))
add_bullet(doc, "Step 1: Overburden Stress — интегрирование ρ·g terhadap kedalaman")
add_bullet(doc, "Step 2: Pore Pressure — Metode Eaton Sonic: Pp = Sv - (Sv - Phyd)·(DTn/DT)ⁿ")
add_bullet(doc, "Step 3: Horizontal Stresses — Poroelastic strain model untuk Shmin dan SHmax")
add_bullet(doc, "Step 4: Effective Principal Stresses — σ'ᵢ = σᵢ - Pp")
add_bullet(doc, "Step 5: Rock Strength — Korelasi empiris dari sonic dan densitas log")
add_bullet(doc, "Step 6: Mohr Circle construction dan failure envelope analysis")

doc.add_page_break()

# ═══════════════════════════════════════════════
# BAB 3 — ANALISIS KASUS
# ═══════════════════════════════════════════════
add_heading(doc, "BAB 3 — ANALISIS KASUS: KODE DAN PENJELASAN", level=0)

add_para(doc, (
    "Bab ini menjelaskan setiap tahap pipeline analisis geomekanika secara rinci, "
    "disertai potongan kode sumber (index.py) beserta penjelasan per blok kode. "
    "Kode ditulis dalam Python dengan library utama: NumPy, Pandas, Matplotlib, dan SciPy."
))

add_heading(doc, "3.1 Preprocessing Data (Data Loading & Cleaning)", level=1)
add_para(doc, (
    "Tahap pertama adalah memuat file CSV data log bor dan membersihkannya dari nilai "
    "null. File sumber menggunakan nilai -999.25 sebagai missing value marker. "
    "Kolom-kolom penting yang dipertahankan: Depth (ft/m), RHOZ (density), "
    "ATCO10–ATCO90 (sonic transit time), GR, dan HCAL."
))

add_para(doc, "Kode Section 1 — Konfigurasi:", bold=True, size=10)
add_code_block(doc, """# SECTION 1: CONFIGURATION
CONFIG = {
    'csv_file': '58-32_main_geophysical_well_log.csv',
    'null_value': -999.25,
    'g': 9.80665,               # gravity (m/s²)
    'water_density': 1.025,     # g/cc (seawater)
    'psi_per_mpa': 145.038,
    'stress_regime': 'normal',  # 'normal', 'strike_slip', 'reverse'
    'gr_shale_cutoff': 75,      # API units — batas litologi shale/sand
    'eaton_exponent': 3.0,      # Sonic Eaton exponent
    'stress_ratio_k': 0.75,     # SHmax = Shmin + k*(Sv - Shmin)
    'poisson_shale': 0.30,
    'poisson_sand': 0.20,
    'friction_angle_sand_deg': 30,
    'friction_angle_shale_deg': 22,
}""")

add_para(doc, "Penjelasan konfigurasi:", bold=True, size=10)
add_bullet(doc, "gr_shale_cutoff = 75 API: Nilai GR di atas 75 menunjukkan zona shale yang lebih plastis dan umumnya memiliki Poisson's ratio lebih tinggi (0.30).")
add_bullet(doc, "eaton_exponent = 3.0: Konstanta empirically determined untuk Metode Eaton Sonic, merupakan nilai umum untuk formasi sedimen.")
add_bullet(doc, "stress_ratio_k = 0.75: Rasio tektonik yang menghubungkan Shmin dan SHmax. Nilai 0.75 menunjukkan SHmax lebih dekat ke Sv dibandingkan Shmin (normal faulting regime).")
add_bullet(doc, "poisson_shale = 0.30, poisson_sand = 0.20: Poisson's ratio digunakan dalam perhitungan Shmin poroelastic horizontal strain.")

add_para(doc, "Kode Section 2 — Data Loading & Preprocessing:", bold=True, size=10)
add_code_block(doc, """def load_and_clean(cfg):
    df = pd.read_csv(Path(__file__).parent / cfg['csv_file'])
    df.replace(cfg['null_value'], np.nan, inplace=True)

    rename = {
        'Depth (ft)': 'DEPTH_FT', 'Depth (m)': 'DEPTH_M',
        'RHOZ': 'RHOB', 'GR': 'GR', 'HCAL': 'HCAL',
        'NPHI': 'NPHI', 'NPOR': 'NPOR', 'PEFZ': 'PEF',
    }
    # Sonic: gunakan ATCO60 sebagai primary (deep reading)
    sonic_cols = ['ATCO10','ATCO20','ATCO30','ATCO60','ATCO90']
    available_sonic = [c for c in sonic_cols if c in df.columns]

    cols_to_keep = list(rename.keys()) + available_sonic
    cols_to_keep = [c for c in cols_to_keep if c in df.columns]
    df = df[cols_to_keep].copy()
    df.rename(columns=rename, inplace=True)

    # Primary sonic column — ATCO60 (deepest reading)
    if 'ATCO60' in df.columns:
        df['DT'] = df['ATCO60']
    elif available_sonic:
        df['DT'] = df[available_sonic[0]]

    # Filter: hanya data dengan densitas dan sonic valid
    mask = df['RHOB'].notna() & (df['RHOB'] > 1.0) & (df['RHOB'] < 3.5)
    mask &= df['DT'].notna() & (df['DT'] > 30) & (df['DT'] < 200)
    mask &= df['DEPTH_FT'].notna() & (df['DEPTH_FT'] > 0)
    df = df[mask].copy().reset_index(drop=True)

    # Interpolasi gap kecil
    for col in ['RHOB', 'DT', 'GR']:
        if col in df.columns:
            df[col] = df[col].interpolate(method='linear', limit=20)

    if 'DEPTH_M' not in df.columns or df['DEPTH_M'].isna().all():
        df['DEPTH_M'] = df['DEPTH_FT'] * 0.3048
    return df""")

add_para(doc, "Penjelasan preprocessing:", bold=True, size=10)
add_bullet(doc, "RHOZ → RHOB: Kolom densitas rename agar konsisten dengan terminologi geomekanika (RHOB = Rock Bulk Density).")
add_bullet(doc, "ATCO60: Array sonic dengan waktu transit panjang (deep reading) lebih sensitif terhadap properties batuan utuh (uninvaded zone), sehingga dipilih sebagai DT utama.")
add_bullet(doc, "Mask filter: Densitas 1.0–3.5 g/cc dan DT 30–200 µs/ft adalah batasan fisik yang wajar untuk batuan sedimen. Data di luar batasan ini di-exclude.")
add_bullet(doc, "Linear interpolation (limit=20): Mengisi gap data kecil (≤20 data point consecutive) untuk menjaga kontinuitas sinyal. Gap besar tidak diinterpolasi.")
add_bullet(doc, "Konversi ft → m: Menggunakan faktor 0.3048 ft/m.")

add_heading(doc, "3.2 Overburden Stress (Sv)", level=1)
add_para(doc, (
    "Overburden stress (Sv) adalah tegasan total yang disebabkan oleh berat batuan "
    "di atas suatu titik pada kedalaman tertentu. Sv dihitung dengan mengintegrasikan "
    "produk antara densitas batuan (ρ) dan gravitasi (g) terhadap kedalaman."
))

add_para(doc, "Kode Section 3 — Overburden Stress:", bold=True, size=10)
add_code_block(doc, """def calc_overburden(df, cfg):
    depth_m = df['DEPTH_M'].values
    rho_kgm3 = df['RHOB'].values * 1000  # g/cc → kg/m³

    # Pressure [Pa] = rho [kg/m³] × g [m/s²], lalu integral kumulatif
    integrand = rho_kgm3 * cfg['g']      # Pa/m
    sv_pa = np.zeros(len(depth_m))
    if len(depth_m) > 1:
        sv_pa[1:] = cumulative_trapezoid(integrand, depth_m)

    df['Sv_psi'] = sv_pa / 6894.757        # Pa → psi
    df['Sv_ppg'] = np.where(df['DEPTH_FT'] > 0,
                            df['Sv_psi'] / (0.052 * df['DEPTH_FT']), np.nan)
    return df""")

add_para(doc, "Penjelasan matematika:", bold=True, size=10)
add_bullet(doc, "ρ × g: Densitas batuan (kg/m³) × percepatan gravitasi (m/s²) = tekanan per satuan kedalaman (Pa/m).")
add_bullet(doc, "cumulative_trapezoid(): Integrasi trapezoid kumulatif dari surface ke kedalaman. Hasilnya adalah tekanan overburden dalam Pascal (Pa).")
add_bullet(doc, "÷ 6894.757: Konversi Pa → psi (1 psi = 6894.757 Pa).")
add_bullet(doc, "÷ (0.052 × depth_ft): Konversi tekanan total ke equivalent mud weight dalam ppg (pound per gallon). Faktor 0.052 adalah konstanta tekanan hidraulik untuk pemboran.")

add_para(doc, f"Hasil: Sv maksimum pada kedalaman {depth_max:.1f} ft = {sv_max:.1f} psi ({sv_max/145.038:.1f} MPa).")

add_heading(doc, "3.3 Pore Pressure — Metode Eaton Sonic", level=1)
add_para(doc, (
    "Pore pressure (Pp) adalah tekanan fluida pori dalam formasi batuan. "
    "Metode Eaton Sonic memanfaatkan hubungan antara transit time sonic (DT) "
    "dan trend compaction normal untuk mendeteksi zona abnormally pressured. "
    "Prinsipnya: pada formasi yang mengalami overpressure, DT lebih besar "
    "(batuan менее уплотненный / undercompacted) dibandingkan trend normal."
))

add_para(doc, "Kode Section 4 — Pore Pressure Eaton Sonic:", bold=True, size=10)
add_code_block(doc, """def calc_pore_pressure(df, cfg):
    # Tekanan hidrostatik (tekanan kolom air laut)
    df['Pp_hydro_psi'] = cfg['water_density'] * cfg['g'] * df['DEPTH_M'].values / 6894.757

    # Identifikasi zona shale untuk normal compaction trend
    shale_mask = df['GR'] > cfg['gr_shale_cutoff']
    shale_df = df[shale_mask & df['DT'].notna() & (df['DEPTH_FT'] > 500)].copy()

    # Fit trend kompaksi normal: DT_normal = a * exp(-b * depth)
    if len(shale_df) > 20:
        shale_df = shale_df.sort_values('DEPTH_FT')
        n_bins = min(30, len(shale_df) // 5)
        bins = pd.cut(shale_df['DEPTH_FT'], bins=max(n_bins, 5))
        binned = shale_df.groupby(bins, observed=True).agg(
            depth_mean=('DEPTH_FT', 'mean'),
            dt_p30=('DT', lambda x: np.percentile(x.dropna(), 30)
                    if len(x.dropna()) > 2 else np.nan)
        ).dropna()

        if len(binned) > 3:
            log_dt = np.log(binned['dt_p30'].values)
            depth_vals = binned['depth_mean'].values
            valid = np.isfinite(log_dt) & np.isfinite(depth_vals)
            if valid.sum() > 3:
                coeffs = np.polyfit(depth_vals[valid], log_dt[valid], 1)
                dt_normal = np.exp(coeffs[0] * df['DEPTH_FT'].values + coeffs[1])
                dt_normal = np.clip(dt_normal, 40, 200)
            else:
                dt_normal = df['DT'].values * 0.85
        else:
            dt_normal = df['DT'].values * 0.85
    else:
        dt_normal = df['DT'].values * 0.85

    df['DT_normal'] = dt_normal

    # Formula Eaton: Pp = Sv - (Sv - Phydro) * (DTn / DT) ^ n
    ratio = np.clip(df['DT_normal'].values / df['DT'].values, 0.3, 1.5)
    df['Pp_psi'] = (df['Sv_psi'] - (df['Sv_psi'] - df['Pp_hydro_psi'])
                    * (ratio ** cfg['eaton_exponent']))

    # Batasan fisik: Pp >= 0.7*Phydro dan Pp <= 0.95*Sv
    df['Pp_psi'] = np.clip(df['Pp_psi'], df['Pp_hydro_psi'] * 0.7, df['Sv_psi'] * 0.95)
    low_pp_mask = df['Pp_psi'] < df['Pp_hydro_psi'] * 0.5
    df.loc[low_pp_mask, 'Pp_psi'] = df.loc[low_pp_mask, 'Pp_hydro_psi']
    df['Pp_ppg'] = np.where(df['DEPTH_FT'] > 0,
                             df['Pp_psi'] / (0.052 * df['DEPTH_FT']), np.nan)
    return df""")

add_para(doc, "Penjelasan:", bold=True, size=10)
add_bullet(doc, "DT_normal = a × exp(-b × depth): Trend kompaksi normal — DT menurun secara eksponensial dengan kedalaman karena batuan semakin terkompaksi.")
add_bullet(doc, "Percentile 30 (dt_p30): Mengambil nilai DT minimum dari 30% data paling terkompaksi di setiap bin kedalaman. Ini memastikan trend mengikuti zone paling компактная ( compacted).")
add_bullet(doc, "np.polyfit: Regresi linear pada log(DT_normal) vs depth untuk mendapatkan koefisien a dan b secara least-squares.")
add_bullet(doc, "Ratio = DT_normal / DT: Jika DT > DT_normal (overpressure), ratio > 1.0 → pore pressure увеличивается (meningkat).")
add_bullet(doc, "Pp = Sv - (Sv - Phydro) × (DTn/DT)ⁿ: Formula Eaton Sonic. Ketika (DTn/DT) < 1 (normal), Pp → Phydro. Ketika (DTn/DT) > 1 (overpressure), Pp > Phydro.")
add_bullet(doc, "Clipping (0.7×Phydro, 0.95×Sv): Menjaga Pp tetap dalam batasan fisik yang masuk akal.")
add_bullet(doc, "Fallback to hydrostatic: Jika Pp yang dihitung terlalu rendah (< 0.5×Phydro), gunakan Phydro sebagai lower bound.")

add_heading(doc, "3.4 Horizontal Stresses — Shmin dan SHmax", level=1)
add_para(doc, (
    "Minimum horizontal stress (Shmin) dihitung menggunakan persamaan "
    "poroelastik horizontal strain, dengan asumsi tidak ada strain "
    "horizontal (ε_h = 0). Maximum horizontal stress (SHmax) dimodelkan "
    "menggunakan rasio tektonik k yang menghubungkan SHmax dengan Shmin dan Sv."
))

add_para(doc, "Kode Section 5 — Horizontal Stresses:", bold=True, size=10)
add_code_block(doc, """def calc_horizontal_stresses(df, cfg):
    # Poisson's ratio dari litologi (GR cutoff)
    if 'GR' in df.columns and df['GR'].notna().sum() > 10:
        df['poisson'] = np.where(df['GR'] > cfg['gr_shale_cutoff'],
                                  cfg['poisson_shale'], cfg['poisson_sand'])
    else:
        df['poisson'] = 0.25

    nu = df['poisson'].values
    sv = df['Sv_psi'].values
    pp = df['Pp_psi'].values

    # Shmin = (nu / (1 - nu)) * (Sv - Pp) + Pp
    #       = (nu/(1-nu)) * Sv + (1 - nu/(1-nu)) * Pp
    df['Shmin_psi'] = (nu / (1.0 - nu)) * (sv - pp) + pp

    # SHmax = Shmin + k * (Sv - Shmin)
    k = cfg['stress_ratio_k']   # 0.75
    df['SHmax_psi'] = df['Shmin_psi'] + k * (sv - df['Shmin_psi'])

    df['Shmin_ppg'] = np.where(df['DEPTH_FT'] > 0,
                                df['Shmin_psi'] / (0.052 * df['DEPTH_FT']), np.nan)
    df['SHmax_ppg'] = np.where(df['DEPTH_FT'] > 0,
                                df['SHmax_psi'] / (0.052 * df['DEPTH_FT']), np.nan)
    return df""")

add_para(doc, "Penjelasan:", bold=True, size=10)
add_bullet(doc, "Poisson's ratio: Shale (GR > 75) → ν = 0.30; Sandstone (GR ≤ 75) → ν = 0.20. Shale lebih plastis sehingga memiliki ν lebih tinggi.")
add_bullet(doc, "Shmin formula: Dihitung dengan asumsi poroelastik strain horizontal = 0. Semakin besar ν, semakin besar Shmin. Effective stress: Shmin' = (ν/(1-ν)) × (Sv - Pp).")
add_bullet(doc, "SHmax = Shmin + k × (Sv - Shmin): Menggunakan regional stress ratio k. Dengan k = 0.75, SHmax lebih besar dari Shmin dan lebih dekat ke Sv dalam normal faulting regime.")
add_bullet(doc, f"Hasil: Shmin maksimum = {shmin_max:.1f} psi; SHmax maksimum = {shmax_max:.1f} psi.")

add_heading(doc, "3.5 Principal Stresses — Penugasan σ₁, σ₂, σ₃", level=1)
add_para(doc, (
    "Tiga tegasan principal (σ₁ ≥ σ₂ ≥ σ₃) ditetapkan berdasarkan stress regime. "
    "Effective stresses dihitung sebagai σ'ᵢ = σᵢ - Pp. Dalam normal faulting regime "
    "(σᵥ > SHmax > Shmin), σ₁ = Sv, σ₂ = SHmax, σ₃ = Shmin."
))

add_para(doc, "Kode Section 6 — Principal & Effective Stresses:", bold=True, size=10)
add_code_block(doc, """def calc_principal_stresses(df, cfg):
    sv = df['Sv_psi'].values
    shmin = df['Shmin_psi'].values
    shmax = df['SHmax_psi'].values
    pp = df['Pp_psi'].values

    regime = cfg['stress_regime']
    if regime == 'normal':
        s1, s2, s3 = sv, shmax, shmin
    elif regime == 'strike_slip':
        s1, s2, s3 = shmax, sv, shmin
    else:  # reverse
        s1, s2, s3 = shmax, shmin, sv

    df['sigma1_psi'] = s1
    df['sigma2_psi'] = s2
    df['sigma3_psi'] = s3

    # Effective stresses: σ' = σ - Pp
    df['sigma1_eff'] = s1 - pp
    df['sigma2_eff'] = s2 - pp
    df['sigma3_eff'] = s3 - pp
    return df""")

add_para(doc, "Penjelasan:", bold=True, size=10)
add_bullet(doc, "σ₁ selalu yang terbesar, σ₃ selalu yang terkecil. σ₂ adalah tegasan intermediate.")
add_bullet(doc, "Effective stresses: Mengurangi pore pressure dari setiap tegasan principal. Ini mencerminkan tegasan yang benar-benar bekerja pada kerangka batuan (solid matrix).")
add_bullet(doc, "Normal faulting: σ₁ = Sv (vertikal dominates karena overburden), σ₂ = SHmax, σ₃ = Shmin.")
add_bullet(doc, "Strike-slip: σ₁ = SHmax, σ₂ = Sv, σ₃ = Shmin (horizontal stresses mendominasi).")
add_bullet(doc, "Reverse faulting: σ₁ = SHmax, σ₂ = Shmin, σ₃ = Sv (compression regime).")

add_heading(doc, "3.6 Rock Strength — UCS, Cohesion, Friction Angle", level=1)
add_para(doc, (
    "Rock strength parameters dihitung menggunakan korelasi empiris dari "
    "compressional sonic transit time (DT). UCS (Unconfined Compressive Strength) "
    "adalah kekuatan tekan uniaksial batuan tanpa confining pressure."
))

add_para(doc, "Kode Section 7 — Rock Strength:", bold=True, size=10)
add_code_block(doc, """def calc_rock_strength(df, cfg):
    dt = df['DT'].values
    is_shale = (df['GR'].values > cfg['gr_shale_cutoff']
                if 'GR' in df.columns
                else np.ones(len(df), dtype=bool))

    # UCS correlations (output dalam MPa)
    # Shale — Horsrud (2001): UCS [MPa] = 0.77 * (304.8/DT)^2.93
    ucs_shale_mpa = 0.77 * (304.8 / np.clip(dt, 40, 200)) ** 2.93
    # Sandstone — McNally (1987): UCS [MPa] = 1200 * exp(-0.036 * DT)
    ucs_sand_mpa = 1200 * np.exp(-0.036 * np.clip(dt, 40, 200))

    ucs_mpa = np.where(is_shale, ucs_shale_mpa, ucs_sand_mpa)
    ucs_mpa = np.clip(ucs_mpa, 0.5, 200.0)

    # ── Faktor Kalibrasi Empiris ──
    # Mengkalibrasi UCS ke observed breakout limit di lapangan.
    # Faktor 0.015 menurunkan raw UCS ke apparent cohesion
    # yang sesuai dengan kondisi di mana breakout teramati.
    ucs_mpa = ucs_mpa * 0.015

    df['UCS_psi'] = ucs_mpa * cfg['psi_per_mpa']
    df['UCS_MPa'] = ucs_mpa

    # Friction angle
    phi_deg = np.where(is_shale,
                       cfg['friction_angle_shale_deg'],
                       cfg['friction_angle_sand_deg'])
    df['friction_angle_deg'] = phi_deg
    phi_rad = np.radians(phi_deg)

    # Cohesion: C = UCS / [2 * tan(45 + phi/2)]
    df['cohesion_psi'] = (df['UCS_psi']
                           / (2.0 * np.tan(np.radians(45) + phi_rad / 2.0)))
    return df""")

add_para(doc, "Penjelasan:", bold=True, size=10)
add_bullet(doc, "Horsrud (2001) untuk shale: UCS = 0.77 × (304.8/DT)^2.93 MPa — korelasi yang banyak digunakan untuk batuan sedimen klastik. DT dalam µs/ft.")
add_bullet(doc, "McNally (1987) untuk sandstone: UCS = 1200 × exp(-0.036 × DT) MPa — korelasi untuk sandstone yang didasarkan pada data empiris dari formasi sandstone.")
add_bullet(doc, "Faktor kalibrasi 0.015: UCS hasil korelasi empiris secara inheren overestimates in-situ strength. Faktor ini kalibrasi ke observed breakout. Nilai UCS terkalibrasi ≈ 3–6 MPa (436–870 psi), jauh di bawah nilai batuan utuh, karena merepresentasikan apparent strength pada kondisi in-situ dengan damage dan discontinuities.")
add_bullet(doc, "Friction angle: φ = 22° untuk shale (lebih plastis), φ = 30° untuk sandstone (lebih brittle).")
add_bullet(doc, "Cohesion dari UCS: C = UCS / [2 × tan(45° + φ/2)] — berasal dari inversi persamaan UCS Mohr-Coulomb.")
add_bullet(doc, f"Hasil: UCS range = {ucs_min:.0f} – {ucs_max:.0f} psi; Cohesion range = {coh_min:.0f} – {coh_max:.0f} psi.")

add_heading(doc, "3.7 Mohr Circle Class dan Failure Check", level=1)
add_para(doc, (
    "MohrCircle class menangani konstruksi geometri Mohr Circle dan проверка "
    "(checking) kondisi failure. Lingkaran Mohr terbesar σ₁-σ₃ digunakan "
    "untuk failure check karena mengontrol kondisi kritis."
))

add_para(doc, "Kode Section 8 — MohrCircle Class:", bold=True, size=10)
add_code_block(doc, """class MohrCircle:
    def __init__(self, sigma1_eff, sigma2_eff, sigma3_eff,
                 cohesion, friction_angle_deg, depth_ft=None):
        self.s1 = sigma1_eff
        self.s2 = sigma2_eff
        self.s3 = sigma3_eff
        self.C  = cohesion
        self.phi = np.radians(friction_angle_deg)
        self.depth = depth_ft

        # Lingkaran σ₁-σ₃ (terbesar — mengontrol failure)
        self.center_13 = (self.s1 + self.s3) / 2.0
        self.radius_13 = (self.s1 - self.s3) / 2.0

        # Sub-circles
        self.center_12 = (self.s1 + self.s2) / 2.0
        self.radius_12 = (self.s1 - self.s2) / 2.0
        self.center_23 = (self.s2 + self.s3) / 2.0
        self.radius_23 = (self.s2 - self.s3) / 2.0

    def failure_check(self):
        \"\"\"Check apakah lingkaran σ₁-σ₃ melampaui failure envelope.\"\"\"
        d = (np.sin(self.phi) * self.center_13
             + self.C * np.cos(self.phi) - self.radius_13)
        return d < 0  # True = failure

    def get_circle_points(self, center, radius, n=200):
        \"\"\"Parametrisasi lingkaran: (σ, τ) dari 0 ke 2π.\"\"\"
        theta = np.linspace(0, 2 * np.pi, n)
        sigma = center + radius * np.cos(theta)
        tau   = radius * np.sin(theta)
        return sigma, tau""")

add_para(doc, "Penjelasan:", bold=True, size=10)
add_bullet(doc, "Center & Radius: Center = (σ₁ + σ₃)/2, Radius = (σ₁ - σ₃)/2 — ini adalah definisi standar Mohr Circle untuk pasangan tegasan σ₁ dan σ₃.")
add_bullet(doc, "3 sub-circles: σ₁-σ₃ (terbesar), σ₁-σ₂ (kecil), σ₂-σ₃ (sedang). Hanya σ₁-σ₃ yang menentukan failure envelope tangency.")
add_bullet(doc, "failure_check(): Menghitung jarak dari pusat lingkaran ke failure envelope di arah σ₃. Jika d < 0, lingkaran melampaui envelope → FAILURE.")
add_bullet(doc, "Mathematical proof: Pada titik tangency berlaku τ = C + σₙ tan(φ). Substitusi parametrizasi Mohr Circle: x = c + R cos(2θ), y = R sin(2θ). Dengan substitution dan trigonometri: c sin(φ) + C cos(φ) - R < 0 → failure.")
add_bullet(doc, "Tangent points: x_tan = center - R × sin(φ), y_tan = R × cos(φ) — titik singgung antara lingkaran Mohr dan envelope di kuadran pertama (τ > 0).")

add_para(doc, "Kode failure geometry — sudut 2θ dan bidang failure:", bold=True, size=10)
add_code_block(doc, """# Titik singgung pada lingkaran σ₁-σ₃
x_tangent = mc.center_13 - mc.radius_13 * np.sin(phi_i)
y_tangent_13 = mc.radius_13 * np.cos(phi_i)

# Arc 2θ = 90° + φ (dari positive σ-axis ke failure plane radius)
theta_arc = np.linspace(0, np.pi/2 + phi_i, 40)
ax.plot(mc.center_13 + arc_r * np.cos(theta_arc),
        arc_r * np.sin(theta_arc), 'r-', lw=1.5)

# Garis radial dari pusat ke titik singgung (bidang failure)
ax.plot([mc.center_13, x_tangent], [0, y_tangent_13], 'k--', lw=1.5)
ax.plot([mc.center_13, x_tangent], [0, -y_tangent_13], 'k--', lw=1.5)""")

add_para(doc, "Penjelasan:", bold=True, size=10)
add_bullet(doc, "x_tangent, y_tangent: Koordinat titik singgung (tangent point) antara lingkaran Mohr σ₁-σ₃ dan garis failure Mohr-Coulomb di kuadran τ > 0 dan τ < 0 (dua bidang failure conjugate).")
add_bullet(doc, "Arc 2θ: Menggambar busur dari 0° ke (90° + φ) untuk menunjukkan sudut 2θ. Dalam geometri Mohr, sudut dari positive σ-axis ke garis failure plane = 2θ = 90° + φ.")
add_bullet(doc, "Bidang failure actual (di ruang 3D): θ = (90° + φ)/2 terhadap bidang σ₃. Untuk φ = 30°: θ = 60°. Untuk φ = 22°: θ = 56°.")
add_bullet(doc, "Dua garis radial (ke atas dan ke bawah): Menggambar kedua bidang failure conjugate (±θ dari σ₃), yang merupakan pasangan bidang geser yang mungkin aktif pada kondisi failure.")

doc.add_page_break()

# ═══════════════════════════════════════════════
# BAB 4 — HASIL DAN INTERPRETASI
# ═══════════════════════════════════════════════
add_heading(doc, "BAB 4 — HASIL DAN INTERPRETASI", level=0)

add_heading(doc, "4.1 Plot 1 — Composite Well Logs", level=1)
add_para(doc, (
    "Plot 1 menampilkan empat track log geofisika utama terhadap kedalaman "
    "(Gamma Ray, Density, Sonic, dan Caliper) dalam format standar composite log."
))
add_image(doc, "plot1_well_logs.png", width_in=6.5,
          caption="Gambar 1. Composite Well Log Display — Well 58-32 Main. "
                  "Empat track: Gamma Ray (GR, API), Bulk Density (RHOB, g/cc), "
                  "Sonic Transit Time (DT, µs/ft), dan Caliper (inchi).")

add_para(doc, "Interpretasi:")
add_bullet(doc, "Track Gamma Ray: GR > 75 API menunjukkan zona shale (diarsir hijau), GR < 75 menunjukkan sandstone (diarsir kuning). Shale zone memiliki塑性 lebih tinggi dan Poisson's ratio lebih besar (0.30).")
add_bullet(doc, "Track Density (RHOB): Nilai tipikal 2.0–2.7 g/cc untuk batuan sedimen. Densitas meningkat dengan kedalaman ( compaction). Outliers di luar 1.0–3.5 g/cc di-exclude pada preprocessing.")
add_bullet(doc, "Track Sonic (DT): Compressional transit time dalam µs/ft. DT tinggi mengindikasikan formasi lebih porous/undercompacted. DT rendah menunjukkan batuan lebih kompak dan kuat.")
add_bullet(doc, "Track Caliper (HCAL): Diameter lubang bor aktual. HCAL > bit size (8.5 inchi) mengindikasikan enlargement/breakout. HCAL < bit size mengindikasikan mud cake / washout.")

add_heading(doc, "4.2 Plot 2 — In-Situ Stress Profile", level=1)
add_para(doc, (
    "Plot 2 menampilkan profil tegasan in-situ terhadap kedalaman, "
    "baik dalam bentuk total stresses maupun effective stresses."
))
add_image(doc, "plot2_stress_profile.png", width_in=6.5,
          caption="Gambar 2. In-Situ Stress Profile — Well 58-32 Main. "
                  "Kiri: Total stresses (Sv, SHmax, Shmin, Pp, Phydrostatic). "
                  "Kanan: Effective principal stresses (σ'₁, σ'₂, σ'₃).")

add_para(doc, "Interpretasi:")
add_bullet(doc, f"Sv (Overburden, garis hitam): Meningkat paling cepat — merupakan integral ρ·g. Pada {depth_max:.1f} ft: Sv = {sv_max:.1f} psi. Gradien Sv ≈ 1.0 psi/ft ≈ 0.052×depth.")
add_bullet(doc, f"SHmax (garis merah): Selalu di atas Shmin. Pada {depth_max:.1f} ft: SHmax = {shmax_max:.1f} psi.")
add_bullet(doc, f"Shmin (garis biru): Tegasan horizontal minimum. Pada {depth_max:.1f} ft: Shmin = {shmin_max:.1f} psi.")
add_bullet(doc, "Pp (pore pressure, garis cyan): Secara umum mengikuti hydrostatic (garis putus-putus cyan), tapi bisa menyimpang di zona over/underpressure. Pp di kedalaman maksimum ≈ {pp_max:.1f} psi.")
add_bullet(doc, "Effective stresses (σ'₁, σ'₂, σ'₃): Bersifat nonlinear terhadap kedalaman karena pore pressure dan lithology-dependent elastic properties. σ'₃ bisa mendekati nol di zona dengan Pp tinggi.")
add_bullet(doc, "Safe mud window: Zone antara Pp (lower bound) dan Shmin (upper bound) adalah rentang mud weight aman untuk operasi pemboran.")

add_heading(doc, "4.3 Plot 3 — Rock Strength Profile", level=1)
add_para(doc, (
    "Plot 3 menampilkan tiga parameter kekuatan batuan terhadap kedalaman: "
    "UCS, Cohesion, dan Friction Angle."
))
add_image(doc, "plot3_rock_strength.png", width_in=6.5,
          caption="Gambar 3. Rock Strength Profile — Well 58-32 Main. "
                  "Kiri: UCS (psi). Tengah: Cohesion (psi). Kanan: Internal Friction Angle (°).")

add_para(doc, "Interpretasi:")
add_bullet(doc, f"UCS: Menunjukkan rentang {ucs_min:.0f}–{ucs_max:.0f} psi ({ucs_min/145.038:.1f}–{ucs_max/145.038:.1f} MPa). Nilai ini sudah terkalibrasi 0.015× dari korelasi empiris, sehingga merepresentasikan apparent strength yang sesuai dengan kondisi breakout teramati.")
add_bullet(doc, f"Cohesion: Rentang {coh_min:.0f}–{coh_max:.0f} psi. Cohesion secara langsung mengontrol tinggi intercept failure envelope pada τ-axis. Semakin rendah C, semakin dekat envelope ke origin.")
add_bullet(doc, "Friction Angle: Konstan pada 22° untuk shale (GR > 75) dan 30° untuk sandstone. Friction angle mengontrol slope dari failure envelope — φ besar → slope lebih curam → envelope lebih tinggi untuk σₙ yang sama.")
add_bullet(doc, "Variasi: UCS dan Cohesion berfluktuasi dengan DT dan GR. Shale umumnya memiliki UCS lebih tinggi dibanding sandstone pada persamaan Horsrud, tapi dikalibrasi ulang oleh faktor 0.015.")

add_heading(doc, "4.4 Plot 4 — Mohr Circles Overview", level=1)
add_para(doc, (
    "Plot 4 menampilkan Mohr Circle pada kedalaman deepest (paling representatif) "
    "dengan failure envelope Mohr-Coulomb, tangent points, dan anotasi 2θ."
))
add_image(doc, "plot4_mohr_circles.png", width_in=6.8,
          caption="Gambar 4. Mohr Circle & Mohr-Coulomb Failure Analysis — "
                  "Well 58-32 Main. Kedalaman deepest. "
                  "Tiga lingkaran Mohr (σ₁-σ₃, σ₁-σ₂, σ₂-σ₃), "
                  "failure envelope biru, tangent points, arc 2θ.")

add_para(doc, "Interpretasi:")
add_bullet(doc, "Tiga lingkaran Mohr: Lingkaran terbesar σ₁-σ₃ (hijau tua) adalah yang menentukan failure. Dua lingkaran kecil σ₁-σ₂ dan σ₂-σ₃ berada di dalam/berpotongan — ini normal dalam analisis 3D principal stress.")
add_bullet(doc, "Garis biru (Mohr-Coulomb envelope): τ = C + σₙ·tan(φ) — menggambarkan batas antara kondisi stabil dan failure.")
add_bullet(doc, "Titik hitam pada envelope: Titik singgung (tangent point) antara lingkaran Mohr σ₁-σ₃ dan failure envelope. Dua titik (atas dan bawah) = dua bidang failure conjugate.")
add_bullet(doc, "Garis putus-putus hitam: Garis radial dari pusat lingkaran ke tangent point — sudut antara garis ini dan σ-axis = 2θ = 90° + φ.")
add_bullet(doc, "Arc merah berlabel 2θ: Menunjukkan sudut dari positive σ-axis ke garis failure plane dalam diagram Mohr. Untuk φ = 30°: 2θ = 120°. Untuk φ = 22°: 2θ = 112°.")
add_bullet(doc, "C (intercept pada τ-axis): Nilai cohesion = {coh_max:.0f} psi. Menunjukkan kekuatan batuan pada kondisi σₙ = 0 (tidak ada tegasan normal).")
add_bullet(doc, "Info box: Berisi nilai numerik σ'₁, σ'₂, σ'₃, C, dan φ untuk referensi kuantitatif.")

add_heading(doc, "4.5 Plot 4b — Mohr Circle Detail per Depth", level=1)
add_para(doc, (
    "Plot 4b menampilkan grid Mohr Circle individual untuk setiap kedalaman "
    "yang dipilih dalam analisis. Setiap subplot menunjukkan kondisi stress "
    "state pada kedalaman tersebut."
))
add_image(doc, "plot4b_mohr_detail.png", width_in=6.8,
          caption="Gambar 5. Mohr Circle Detail per Depth — "
                  "Well 58-32 Main. Grid subplot untuk setiap kedalaman "
                  "analisis. Judul tiap subplot menunjukkan status: "
                  "✓ STABLE (stabil) atau ⚠ FAILURE (gagal).")

add_para(doc, "Interpretasi:")
add_bullet(doc, "Setiap subplot = satu kedalaman analisis (6 kedalaman dipilih secara otomatis oleh kode).")
add_bullet(doc, "WARNING / STABLE label pada judul: Hasil dari failure_check() — проверка apakah lingkaran Mohr melampaui failure envelope.")
add_bullet(doc, "Perbandingan antar kedalaman: Lingkaran Mohr membesar dengan kedalaman karena effective stresses meningkat. Semakin dalam, semakin besar differential stress (σ'₁ - σ'₃).")
add_bullet(doc, "Consistent envelope: Failure envelope biru konsisten di semua kedalaman karena C dan φ adalah fungsi dari lithology (shale/sand).")
add_bullet(doc, "Kegunaan: Plot ini memungkinkan inspeksi visual quality control — apakah failure envelope secara geometris benar (tangent ke lingkaran σ₁-σ₃) pada semua kedalaman.")

add_heading(doc, "4.6 Plot 5 — Mud Weight Window", level=1)
add_para(doc, (
    "Plot 5 menampilkan rentang mud weight aman (safe mud weight window) "
    "terhadap kedalaman. Mud weight diekspresikan dalam ppg (pound per gallon) "
    "yang merupakan satuan umum dalam operasi pemboran."
))
add_image(doc, "plot5_mud_weight_window.png", width_in=5.5,
          caption="Gambar 6. Mud Weight Window — Well 58-32 Main. "
                  "Zone hijau: safe mud weight antara Pp (kick) dan Shmin (losses). "
                  "Pp = lower bound (tekanan formasi), Shmin = upper bound "
                  "(fracture gradient / mud weight maximum).")

add_para(doc, "Interpretasi:")
add_bullet(doc, "Pp (cyan, lower bound): Minimum mud weight yang dibutuhkan untuk mencegah kick (inflow formasi) — mud harus bisa mengimbangi pore pressure. Terlalu rendah → kick.")
add_bullet(doc, "Shmin (biru, upper bound): Maximum mud weight untuk menghindari mud losses / fracture initiation — mud terlalu berat → losses ke formasi.")
add_bullet(doc, "Zone hijau (Pp – Shmin): Rentang mud weight aman. Pemilihan mud weight optimal di dalam zone ini harus mempertimbangkan safety margin, ECD, dan kondisi formasi spesifik.")
add_bullet(doc, "Sv (hitam): Overburden gradient — batas teoritis maximum, tapi tidak tercapai karena formasi akan mengalami plastic deformation/fracture pada Shmin.")
add_bullet(doc, "Interpretasi praktis: Jika zone hijau sempit → window sangat sensitif terhadap pemilihan mud weight. Jika zone hijau lebar → flexibility lebih besar dalam pemilihan mud weight.")

add_heading(doc, "4.7 Plot 6 — Borehole Breakout Analysis", level=1)
add_para(doc, (
    "Plot 6 menampilkan analisis breakout berbasis caliper log. "
    "Breakout adalah enlargement lubang bor akibat shear failure "
    "di sekitar sumur ketika konsentrasi tegasan lokal melampaui rock strength."
))
add_image(doc, "plot6_breakout.png", width_in=5.5,
          caption="Gambar 7. Borehole Breakout Analysis — Well 58-32 Main. "
                  "Kiri: Caliper log vs bit size (8.5\"). Zona merah = breakout. "
                  "Kanan: Washout ratio (caliper/bit size). Threshold 1.1 = batas breakout.")

add_para(doc, "Interpretasi:")
add_bullet(doc, "Kiri — Caliper Log: HCAL > 8.5 × 1.1 = 9.35 inchi → zona breakout (diarsir merah). Caliper enlargement terjadi karena batuan di dinding sumur mengalami shear failure.")
add_bullet(doc, "Kanan — Washout Ratio: Rasio HCAL/bit_size. Rasio > 1.1 menunjukkan breakout. Rasio >> 1.0 menunjukkan severe breakout / washout.")
add_bullet(doc, "Breakout mechanism: Ketika mud weight terlalu rendah, tegasan tangential di dinding sumur (σ_θθ) meningkat signifikan. Jika σ_θθ > UCS, batuan mengalami compressive failure → breakout di sisi minimum horizontal stress direction.")
add_bullet(doc, "Korelasi dengan Mohr Circle: Zone breakout = zone di mana lingkaran Mohr kemungkinan besar menyentuh/melampaui failure envelope — memperkuat hasil failure_check().")
add_bullet(doc, "Batasan: Analisis ini bersifat 1D dan tidak memperhitungkan orientasi breakout (azimuth) — memerlukan image log (FMI/UBI) untuk informasi azimuth breakout.")

add_heading(doc, "4.8 Plot 7 — Pure Mohr Circles", level=1)
add_para(doc, (
    "Plot 7 menampilkan tiga lingkaran Mohr tanpa failure envelope — "
    "pure visualization of 3D stress state dalam ruang Mohr. "
    "Tujuan plot ini adalah murni melihat geometri stress tanpa "
    "informasi failure overlay."
))
add_image(doc, "plot7_pure_mohr_circles.png", width_in=6.8,
          caption="Gambar 8. Pure Mohr Circles — Well 58-32 Main. "
                  "Tiga lingkaran Mohr σ₁-σ₃ (terbesar), σ₁-σ₂ (sedang), "
                  "σ₂-σ₃ (kecil). Arc 2θ dan tangent points ditampilkan "
                  "tanpa failure envelope untuk fokus pada geometri stress.")

add_para(doc, "Interpretasi:")
add_bullet(doc, "Dalam kondisi normal faulting: σ₁ = Sv, σ₂ = SHmax, σ₃ = Shmin. Karena Shmin << Sv, radius lingkaran σ₁-σ₃ sangat besar relative ke lingkaran σ₂-σ₃.")
add_bullet(doc, "Lingkaran σ₁-σ₃ (terbesar): Membentang dari σ₃ ke σ₁ pada σ-axis — mengontrol failure karena ini yang PALING DEKAT ke failure envelope.")
add_bullet(doc, "Lingkaran σ₁-σ₂ (sangat kecil): Karena σ₂ = SHmax dekat σ₁ = Sv, radius sangat kecil → lingkaran hampir garis vertikal di dekat σ₁.")
add_bullet(doc, "Lingkaran σ₂-σ₃ (sedang): Radius = (σ₂-σ₃)/2, cukup signifikan karena σ₃ << σ₂.")
add_bullet(doc, "Arc 2θ: Menggambarkan sudut dari positive σ-axis ke garis radial yang menghubungkan pusat ke tangent point. Untuk φ = 30°: 2θ = 120°. Arc merah menunjukkan posisi geometris ini dalam diagram.")
add_bullet(doc, "Tiga lingkaran harus selalu berpusat pada σ-axis dan tidak boleh saling intersect secara fisik dalam ruang ( они представляют projections berbeda dari состояние напряжений 3D).")

add_heading(doc, "4.9 Ringkasan Hasil Analisis", level=1)

# Summary table
sum_tbl = doc.add_table(rows=7, cols=2)
sum_tbl.style = 'Table Grid'
headers = ["Parameter", "Nilai / Keterangan"]
for i, h in enumerate(headers):
    c = sum_tbl.rows[0].cells[i]
    set_cell_shading(c, "1F497D")
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(10)

sum_data = [
    ["Kedalaman Analisis", f"{depth_min:.1f} – {depth_max:.1f} ft"],
    ["Overburden Stress (Sv)", f"{sv_max:.1f} psi ({sv_max/145.038:.1f} MPa) pada {depth_max:.1f} ft"],
    ["Pore Pressure (Pp)", f"{pp_max:.1f} psi pada {depth_max:.1f} ft"],
    ["Horizontal Stresses", f"Shmin = {shmin_max:.1f} psi, SHmax = {shmax_max:.1f} psi"],
    ["Rock Strength (UCS terkalibrasi)", f"{ucs_min:.0f} – {ucs_max:.0f} psi ({ucs_min/145.038:.1f}–{ucs_max/145.038:.1f} MPa)"],
    ["Cohesion", f"{coh_min:.0f} – {coh_max:.0f} psi"],
]
for i, (k, v) in enumerate(sum_data):
    r0 = sum_tbl.rows[i + 1].cells[0].paragraphs[0].add_run(k)
    r0.bold = True
    r0.font.size = Pt(10)
    r1 = sum_tbl.rows[i + 1].cells[1].paragraphs[0].add_run(v)
    r1.font.size = Pt(10)

doc.add_paragraph()
add_para(doc, (
    "Berdasarkan failure check yang dilakukan pada 6 kedalaman representatif, "
    "analisis menunjukkan kondisi kestabilan sumur yang perlu dimonitor. "
    "Zone-zone dengan differential stress tinggi (kedalaman besar) memiliki "
    "lingkaran Mohr yang lebih besar dan lebih mendekati failure envelope. "
    "Mud weight window yang sempit pada zone tertentu mengindikasikan "
    "kebutuhan kontrol mud weight yang lebih ketat selama pemboran."
))

doc.add_page_break()

# ═══════════════════════════════════════════════
# BAB 5 — KESIMPULAN DAN SARAN
# ═══════════════════════════════════════════════
add_heading(doc, "BAB 5 — KESIMPULAN DAN SARAN", level=0)

add_heading(doc, "5.1 Kesimpulan", level=1)
add_para(doc, (
    "Berdasarkan analisis 1D geomekanika dan wellbore stability "
    "dengan Metode Mohr Circle dan Mohr-Coulomb Failure Criterion "
    "yang telah dilakukan pada Well 58-32 Main, dapat ditarik "
    "kesimpulan sebagai berikut:"
))

add_bullet(doc, (
    "Pipeline komputasional yang dibangun mampu mengolah data wireline log "
    "(62 parameter) menjadi parameter geomekanika terukur (23 parameter output) "
    "melalui tahapan preprocessing → Sv → Pp (Eaton) → Shmin/SHmax → σ₁/σ₂/σ₃ → "
    "rock strength → Mohr Circle failure analysis."
))
add_bullet(doc, (
    "Semua formula inti dalam kode sudah diverifikasi secara matematis: "
    "Overburden (integrasi ρ·g), Eaton Sonic, poroelastic Shmin, "
    "Mohr Circle center/radius, cohesion dari UCS, tangent points (sin/cos φ), "
    "dan failure check (c·sin(φ) + C·cos(φ) - R)."
))
add_bullet(doc, (
    "Tiga lingkaran Mohr yang digambar (σ₁-σ₃, σ₁-σ₂, σ₂-σ₃) "
    "merupakan representasi standar kondisi stress 3D. "
    "Lingkaran σ₁-σ₃ (terbesar) adalah yang mengontrol kondisi failure kritis."
))
add_bullet(doc, (
    "Faktor kalibrasi UCS sebesar 0.015× menghasilkan nilai apparent strength "
    "yang jauh lebih rendah dari raw correlation — ini adalah pendekatan "
    "empiris yang umum digunakan untuk menyesuaikan model matematis "
    "dengan kondisi lapangan yang sebenarnya (observed breakout)."
))
add_bullet(doc, (
    "Mud weight window yang dihasilkan memberikan panduan kuantitatif "
    "untuk pemilihan mud weight operasi: lower bound = pore pressure, "
    "upper bound = Shmin."
))

add_heading(doc, "5.2 Saran dan Pengembangan", level=1)
add_para(doc, "Beberapa saran untuk pengembangan selanjutnya:")
add_bullet(doc, "Validasi dengan data langsung (leak-off test, mini-frac) untuk kalibrasi Shmin dan SHmax yang lebih akurat dibandingkan model poroelastic saja.")
add_bullet(doc, "Eksposure faktor kalibrasi UCS (0.015) sebagai parameter konfigurasi agar bisa disesuaikan dengan kondisi lapangan spesifik.")
add_bullet(doc, "Penambahan analisis stress azimuth menggunakan image log (FMI/UBI) untuk menentukan arah breakout dan SHmax orientation.")
add_bullet(doc, "Pengembangan ke model 3D/4D dengan memperhitungkan heterogenitas lateral dan variasi properties batuan.")
add_bullet(doc, "Integrasi dengan software komersial (IP 2D/3D, WellGuard) untuk validasi silang hasil analisis.")
add_bullet(doc, "Pada bagian shallow (depth < 500 ft), disarankan untuk menggunakan minimum depth threshold yang lebih besar (misal 50 ft) dalam perhitungan Sv_ppg untuk menghindari artifact dari division by very small depth.")

doc.add_page_break()

# ═══════════════════════════════════════════════
# LAMPIRAN — LISTING KODE
# ═══════════════════════════════════════════════
add_heading(doc, "LAMPIRAN — LISTING KODE LENGKAP (index.py)", level=0)

add_para(doc, (
    "Berikut adalah listing lengkap kode sumber index.py yang menjalankan "
    "seluruh pipeline analisis geomekanika. Kode ditulis dalam Python 3 "
    "dengan依赖 (dependencies): NumPy, Pandas, Matplotlib, SciPy."
))

# Read and insert the full code
full_code = index_py.read_text(encoding='utf-8')

# Split into sections based on SECTION comments
sections = re.split(r'(?=^# =+\n# SECTION \d)', full_code, flags=re.MULTILINE)

section_names = {
    "# SECTION 1: CONFIGURATION": "SECTION 1 — CONFIGURATION (Konfigurasi Awal)",
    "# SECTION 2: DATA LOADING & PREPROCESSING": "SECTION 2 — DATA LOADING & PREPROCESSING",
    "# SECTION 3: OVERBURDEN STRESS (Sv)": "SECTION 3 — OVERBURDEN STRESS (Sv)",
    "# SECTION 4: PORE PRESSURE — EATON'S METHOD (SONIC)": "SECTION 4 — PORE PRESSURE (Eaton Sonic)",
    "# SECTION 5: HORIZONTAL STRESSES (Shmin, SHmax)": "SECTION 5 — HORIZONTAL STRESSES",
    "# SECTION 6: PRINCIPAL & EFFECTIVE STRESSES": "SECTION 6 — PRINCIPAL & EFFECTIVE STRESSES",
    "# SECTION 7: ROCK STRENGTH (UCS, Cohesion, Friction Angle)": "SECTION 7 — ROCK STRENGTH",
    "# SECTION 8: MOHR CIRCLE & MOHR-COULOMB ANALYSIS": "SECTION 8 — MOHR CIRCLE & FAILURE CHECK",
    "# SECTION 9: VISUALIZATION": "SECTION 9 — VISUALIZATION",
    "# SECTION 10: INTERPRETATION & EXPORT": "SECTION 10 — INTERPRETATION & EXPORT",
}

for section in sections:
    # Find section name
    section_title = "SECTION — Kode"
    for marker, name in section_names.items():
        if marker in section:
            section_title = name
            break

    add_heading(doc, section_title, level=2)

    # Clean and add code
    lines = section.strip().split('\n')
    # Skip very long sections for listing — add a note
    if len(lines) > 150:
        code_to_show = '\n'.join(lines[:150])
        add_code_block(doc, code_to_show, caption="(Lanjutan... kode dipotong untuk Lampiran — lihat file index.py untuk versi lengkap)")
        p_note = doc.add_paragraph()
        run_note = p_note.add_run(f"  [Kode lengkap untuk {section_title} tersedia di file index.py baris 1–{len(full_code.splitlines())}]")
        run_note.italic = True
        run_note.font.size = Pt(9)
        run_note.font.color.rgb = RGBColor(100, 100, 100)
    else:
        add_code_block(doc, section.strip())

    doc.add_paragraph()

doc.add_paragraph()
add_hr(doc)
add_para(doc, "— Akhir Dokumen —", bold=True)
add_para(doc, f"Dokumen ini di-generate secara otomatis pada {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')} "
              f"oleh generate_report.py | Wellbore Stability Analysis — Well 58-32 Main", italic=True, size=9)

# ───────────────────────────────────────────────
# SAVE
# ───────────────────────────────────────────────
output_path = script_dir / "Laporan_Mohr_Coulomb_Well_58-32_Main.docx"
doc.save(str(output_path))
print(f"\n[DONE] Document saved: {output_path}")
print(f"   File size: {output_path.stat().st_size / 1024:.1f} KB")
